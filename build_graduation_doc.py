from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"C:\other\新建文件夹\毕设\基于大模型（LLM）的智能穿搭推荐")
OUT = ROOT / "基于大语言模型的智能穿搭推荐系统_毕业设计成果.docx"
ASSETS = ROOT / ".codex_tmp" / "graduation-assets"
ASSETS.mkdir(parents=True, exist_ok=True)
FONT = r"C:\Windows\Fonts\msyh.ttc"

def f(size, bold=False):
    return ImageFont.truetype(FONT, size, index=0)

def editorial_f(size, active=False):
    return ImageFont.truetype(r"C:\Windows\Fonts\STZHONGS.TTF" if active else r"C:\Windows\Fonts\STSONG.TTF", size)

def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def border_cell(cell, color="A6A6A6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)

def set_cell(cell, text, bold=False, center=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, "宋体", size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, "F2F4F7" if bold else "FFFFFF")
    border_cell(cell)

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    elem = OxmlElement("w:tblHeader")
    elem.set(qn("w:val"), "true")
    tr_pr.append(elem)

def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Cm(widths[i])

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        set_cell(table.rows[0].cells[i], value, bold=True, center=True, size=8.5)
    set_repeat_table_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell(cells[i], value, center=i == 0, size=8.5)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph()
    return table

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, "宋体", 9)

def text(doc, value, indent=True):
    p = doc.add_paragraph()
    p.style = "BodyText"
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(value)
    set_font(r, "宋体", 10.5)
    return p

def bullet(doc, value):
    p = doc.add_paragraph(style="BodyText")
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.37)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run("• " + value)
    set_font(r, "宋体", 10.5)

def heading(doc, value, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(13 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(value)
    set_font(r, "黑体", 16 if level == 1 else (14 if level == 2 else 12), bold=True)
    return p

def add_page_field(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

def add_toc_field(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(fld)

def diagram_flow(path):
    im = Image.new("RGB", (1800, 620), "white")
    d = ImageDraw.Draw(im)
    nodes = [(70,240,300,370,"登录与偏好\n设置"),(390,240,650,370,"场景输入\n天气/场合/风格"),(740,240,1000,370,"衣橱检索与\n规则过滤"),(1090,240,1350,370,"LLM 生成\n搭配方案"),(1440,240,1730,370,"展示、收藏\n与反馈")]
    for x1,y1,x2,y2,label in nodes:
        d.rounded_rectangle((x1,y1,x2,y2), radius=18, fill="#F5F7FA", outline="#34495E", width=4)
        for j, line in enumerate(label.split("\n")):
            box=d.textbbox((0,0),line,font=f(31)); w=box[2]-box[0]
            d.text(((x1+x2-w)//2, y1+30+j*42), line, fill="#17212B", font=f(31))
    for a,b in zip(nodes,nodes[1:]):
        d.line((a[2]+8,305,b[0]-15,305), fill="#557A95", width=5)
        d.polygon([(b[0]-15,305),(b[0]-35,293),(b[0]-35,317)], fill="#557A95")
    d.text((70,65), "智能穿搭推荐核心业务流程", fill="#17212B", font=f(48))
    d.text((70,135), "结构化输入 -> 可解释检索 -> 受约束的模型生成 -> 用户反馈闭环", fill="#52616B", font=f(25))
    im.save(path)

def diagram_arch(path):
    im = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(im)
    d.text((70,45), "系统总体架构", fill="#17212B", font=f(48))
    layers=[("表现层",["Vue 3 + Vite","Element Plus / Pinia","放射式功能切换"],"#EEF4F7"),("业务服务层",["Spring Boot 3","Spring Security + JWT","MyBatis-Plus / 业务编排"],"#F6F2EA"),("智能推荐层",["场景解析与提示词模板","衣橱/向量检索","Qwen 兼容 LLM 接口"],"#F2F0F6"),("数据与基础设施层",["PostgreSQL + pgvector","Redis 缓存","对象存储 / 日志监控"],"#EEF5EF")]
    y=145
    for title,items,color in layers:
        d.rounded_rectangle((90,y,1710,y+170), radius=18, fill=color, outline="#596D7C", width=3)
        d.text((125,y+60),title,fill="#17212B",font=f(31))
        x=410
        for item in items:
            d.rounded_rectangle((x,y+42,x+360,y+128),radius=12,fill="white",outline="#9AA7B1",width=2)
            box=d.textbbox((0,0),item,font=f(23)); w=box[2]-box[0]
            d.text((x+(360-w)//2,y+71),item,fill="#273746",font=f(23))
            x+=405
        y+=210
    im.save(path)

def diagram_er(path):
    im=Image.new("RGB",(1800,980),"white"); d=ImageDraw.Draw(im)
    d.text((65,42),"核心数据实体关系",fill="#17212B",font=f(48))
    boxes=[(90,190,390,365,"sys_user\n用户与账户"),(525,155,850,385,"garment\n单件衣物"),(525,555,850,785,"outfit_plan\n搭配方案"),(1010,155,1350,385,"recommendation_session\n推荐会话"),(1010,555,1350,785,"outfit_item\n方案明细"),(1480,350,1740,580,"favorite_outfit\n收藏与反馈")]
    for x1,y1,x2,y2,label in boxes:
        d.rounded_rectangle((x1,y1,x2,y2),radius=15,fill="#F7F8FA",outline="#455A64",width=3)
        for j,line in enumerate(label.split("\n")):
            bb=d.textbbox((0,0),line,font=f(27)); d.text(((x1+x2-(bb[2]-bb[0]))//2,y1+55+j*43),line,fill="#17212B",font=f(27))
    links=[((390,278),(525,270),"1:N"),((390,315),(525,670),"1:N"),((850,270),(1010,270),"N:1"),((850,670),(1010,670),"1:N"),((1350,270),(1480,430),"1:N"),((1350,670),(1480,500),"N:1")]
    for (a,b,label) in links:
        d.line((*a,*b),fill="#617D8A",width=4); d.text(((a[0]+b[0])//2-20,(a[1]+b[1])//2-40),label,fill="#617D8A",font=f(20))
    im.save(path)

def diagram_ui(path):
    im=Image.new("RGB",(1800,950),"#FBFAF8"); d=ImageDraw.Draw(im)
    d.text((75,55),"知己 / 钟表式功能切换",fill="#17212B",font=f(46))
    # Labels are fixed like clock-face markers. Only the two hands rotate with the active route.
    cx,cy=865,715
    labels=[("知己",805,165,"#D2CEC8",False),("衣橱",320,370,"#D2CEC8",False),("风潮",175,655,"#D2CEC8",False),("档案",1190,655,"#D2CEC8",False),("搭配",1195,360,"#1A1715",True),("设定",835,760,"#D2CEC8",False)]
    for lab,x,y,col,active in labels:
        d.text((x,y),lab,fill=col,font=editorial_f(56 if active else 46, active))
    d.line((cx,cy,1085,475),fill="#9CA0A1",width=4)
    d.line((cx,cy,1360,390),fill="#9CA0A1",width=4)
    d.ellipse((cx-10,cy-10,cx+10,cy+10),fill="#333333")
    im.save(path)

def build():
    for fn, maker in [("flow.png",diagram_flow),("architecture.png",diagram_arch),("er.png",diagram_er),("ui.png",diagram_ui)]: maker(ASSETS/fn)
    doc=Document()
    sec=doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.left_margin=sec.right_margin=Cm(2.5); sec.top_margin=Cm(3.2); sec.bottom_margin=Cm(2.5)
    styles=doc.styles
    styles["Normal"].font.name="宋体"; styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); styles["Normal"].font.size=Pt(10.5)
    if "BodyText" not in styles: styles.add_style("BodyText",1)
    styles["BodyText"].font.name="宋体"; styles["BodyText"]._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体")
    for name in ("Heading 1","Heading 2","Heading 3"):
        styles[name].font.name="黑体"; styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体")
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(footer.add_run("第 "),"宋体",9); add_page_field(footer); set_font(footer.add_run(" 页"),"宋体",9)
    # Cover
    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("湖南科技职业学院"); set_font(r,"黑体",24,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("2026届学生毕业设计成果"); set_font(r,"黑体",25,True)
    doc.add_paragraph(); doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("知己 - 基于大语言模型的智能穿搭推荐系统"); set_font(r,"黑体",23,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("（企业软件方向）"); set_font(r,"宋体",14)
    for _ in range(4): doc.add_paragraph()
    meta=[("学生姓名","【待填写】"),("学    号","【待填写】"),("二级学院","软件学院"),("专业班级","【待填写】"),("指导教师","【待填写】"),("完成日期","2026年  月  日")]
    t=doc.add_table(rows=len(meta),cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; set_table_widths(t,[4,9])
    for row,(k,v) in zip(t.rows,meta): set_cell(row.cells[0],k,True,True,11); set_cell(row.cells[1],v,False,True,11)
    doc.add_page_break()
    # statement
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("毕业设计真实性承诺及指导教师声明"); set_font(r,"黑体",18,True)
    text(doc,"本人郑重声明：本毕业设计在指导教师指导下独立完成。文档中的技术路线、设计说明、引用资料均应在提交前由本人结合实际实现与测试记录复核；未注明的内容不应作为已完成事实使用。",False)
    text(doc,"学生签名：____________________      指导教师签名：____________________",False)
    text(doc,"日期：2026年____月____日",False)
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("目  录"); set_font(r,"黑体",18,True)
    p=doc.add_paragraph(); add_toc_field(p)
    heading(doc,"1 设计任务",1)
    heading(doc,"1.1 项目背景",2)
    text(doc,"服饰搭配通常需要同时考虑天气、场合、个人风格、已有衣物和颜色协调等因素。传统电商推荐多以商品相似度或热门程度为依据，难以围绕用户已有衣橱给出可执行的成套方案。大语言模型具备自然语言理解、约束整合与解释生成能力，可将用户的场景描述转化为结构化搭配条件，并输出包含单品、配色、理由和替代方案的推荐结果。")
    text(doc,"本项目拟设计并实现一个面向普通用户的智能穿搭推荐系统。系统以用户衣橱为主要数据来源，结合天气和场合信息，通过检索、规则校验与大语言模型协同生成搭配方案。项目的根本目标不是生成泛化的时尚文案，而是在用户可选衣物范围内给出可追溯、可修改、可收藏的推荐。")
    heading(doc,"1.2 项目范围",2)
    text(doc,"系统面向注册用户与内容管理员两类角色。前台围绕一次完整的“录入衣物 - 描述场景 - 获得方案 - 反馈优化”闭环设计，管理员负责管理服饰标签、场景规则和异常推荐记录。第一期不包含在线交易、虚拟试衣和社交发布等高成本功能。")
    for title,body in [
        ("1.2.1 用户认证与偏好管理","支持账号注册、登录、JWT 会话校验和个人偏好维护。偏好包含性别展示选项、常用场合、色彩倾向、预算敏感度和风格标签；敏感认证信息采用不可逆密码散列保存。"),
        ("1.2.2 我的衣橱管理","用户可上传衣物图片，维护品类、颜色、季节、厚薄、风格和可穿状态等标签，并按品类、颜色和季节筛选。图片文件仅保存对象存储地址，业务数据库保存元数据。"),
        ("1.2.3 智能搭配推荐","用户输入出行场景、天气、温度和额外要求后，系统从个人衣橱检索候选衣物，使用规则先过滤明显不适配组合，再请求大语言模型生成结构化方案。结果包括上装、下装、鞋履、配饰、配色建议和推荐理由。"),
        ("1.2.4 搭配收藏与反馈","用户可收藏方案、标注满意度、替换某件单品并查看历史推荐。反馈将用于调整后续排序，不直接把用户原始文本作为模型长期训练数据。"),
        ("1.2.5 内容与规则管理","管理员维护服饰标签、场景字典、敏感词和基础搭配规则，查看模型调用异常与推荐质量反馈，为迭代提供可审计依据。"),
        ("1.2.6 风潮趋势聚合","风潮页面以抖音为主来源，按配置接入小红书、微博、Instagram 等补充来源的公开时尚话题和趋势摘要。第一期仅聚合标题或话题、公开热度、原始链接和采集时间；后端仅通过平台允许的 API、公开页面或授权数据源获取信息，不绕过登录、验证码、访问限制或反爬措施。完整性以已许可来源的公开趋势元数据为边界，不承诺采集平台全站内容。"),
        ("1.2.7 个人风格分析","用户完成基础信息、风格偏好、色彩倾向和衣橱标签后，系统调用阿里云百炼平台的通义千问 API 生成结构化个人风格档案。档案包含推荐风格、可尝试风格、适配颜色、推荐单品和简短理由；风潮页面读取该档案并与全局趋势分区展示，不将全局热度直接视为个人偏好。")]:
        heading(doc,title,3); text(doc,body)
    heading(doc,"1.3 项目风险分析",2)
    text(doc,"项目风险主要集中在模型输出不稳定、衣物标签质量、第三方模型接口可用性以及个人信息保护。系统采用“结构化输出约束 + 业务规则二次校验 + 降级提示”的方式控制风险。")
    caption(doc,"表1.1 项目风险分析表")
    add_table(doc,["序号","风险","概率","影响","应对措施"],[
        ["1","模型返回非 JSON 或遗漏单品","中","高","限定 JSON Schema，解析失败时重试一次并返回可理解的降级提示。"],
        ["2","衣橱为空或标签缺失","高","中","引导补充最少 3 件基础单品；对缺失字段采用明确的未标注状态。"],
        ["3","LLM 接口超时或限流","中","高","设置超时、熔断和缓存最近方案；提示用户稍后重试。"],
        ["4","图片与隐私数据泄露","低","高","鉴权访问、对象存储私有读、最小化日志和定期清理无效文件。"],
        ["5","功能范围扩张","中","中","以衣橱和推荐闭环为验收边界，虚拟试衣等功能列为后续迭代。"],
        ["6","第三方趋势来源不可用或限制访问","中","中","按来源独立降级，展示最后一次成功快照与采集时间，不规避平台访问限制。"]],[1,4,1.5,1.5,8])
    heading(doc,"1.4 任务分配",2)
    text(doc,"本项目按个人毕业设计组织。需求分析、原型设计、前后端开发、测试与文档由学生独立完成；指导教师负责技术路线与阶段成果审核。")
    caption(doc,"表1.2 任务分配表")
    add_table(doc,["阶段","主要任务","责任人","交付物"],[
        ["需求分析","用户场景、范围、用例和原型","学生","需求与设计说明"],["概要设计","架构、数据库、接口与提示词约束","学生","技术方案"],["详细实现","Java 后端、Vue 3 前端、模型接入","学生","可运行系统"],["测试验收","单元、接口和核心链路测试","学生","测试记录"],["过程指导","方案审阅与质量把关","指导教师","指导意见"]],[2.2,5.5,2.5,4.2])
    heading(doc,"1.5 项目所需资源",2)
    caption(doc,"表1.3 资源需求表")
    add_table(doc,["资源","版本或规格","用途"],[
        ["JDK","17 LTS","编译和运行 Spring Boot 服务"],["Node.js","20 LTS","Vue 3 前端构建"],["PostgreSQL","16 + pgvector","业务数据与向量检索"],["Redis","3.2.100（本机）；云端使用 Redis 7+","天气、趋势快照和限流缓存"],["Docker Compose","v2","云服务器一致化部署"],["Jsoup","MIT 开源 HTML 解析库","解析允许采集的公开网页；不用于规避访问限制"],["Open-Meteo","开源天气 API，非商业免费使用","按城市获取实时天气，辅助搭配规则过滤"],["阿里云百炼通义千问 API","由 DASHSCOPE_API_KEY 配置","个人风格分析与搭配文本生成"]],[3.3,4.2,6.9])
    heading(doc,"2 设计思路与技术方案",1)
    text(doc,"系统采用前后端分离架构。前端负责场景表单、衣橱维护和结果呈现；后端负责鉴权、数据管理、检索编排、规则校验和模型调用。模型只输出受约束的推荐候选，最终展示的方案必须通过库存、季节、颜色和用户衣橱归属校验。")
    heading(doc,"2.1 系统核心业务流程图",2)
    doc.add_picture(str(ASSETS/"flow.png"),width=Cm(16)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; caption(doc,"图2.1 智能搭配推荐核心业务流程图")
    heading(doc,"2.2 系统用例图",2)
    text(doc,"普通用户可注册登录、维护衣橱、提交场景、获取推荐、收藏方案和提交反馈；管理员可维护标签与规则、查看推荐异常。推荐用例依赖用户鉴权、衣橱候选检索和模型服务三个前置能力。")
    caption(doc,"表2.1 系统主要用例")
    add_table(doc,["角色","主要用例","说明"],[
        ["普通用户","注册登录、偏好设置、衣橱管理","维护个人可用于推荐的衣物与偏好。"],["普通用户","场景推荐、替换单品、收藏反馈","获取并调整搭配方案，形成个人历史。"],["管理员","标签管理、规则管理、异常查看","保证字典与规则可维护、模型异常可追踪。"]],[2.2,5.2,7])
    heading(doc,"2.3 用例分析",2)
    for index,(code,name,desc,pre,flow,post) in enumerate([
        ("UC-01","用户登录认证","用户使用账号和密码进入系统并获得访问令牌。","用户已注册且账号未被禁用。","输入账号密码；后端校验密码散列；生成 JWT；前端保存短期访问令牌。","用户进入首页，可访问自己的衣橱和历史方案。"),
        ("UC-02","新增衣物","用户上传衣物图片并维护属性，形成推荐候选。","用户已登录，图片格式符合限制。","选择图片；填写品类、颜色、季节等标签；后端校验；保存图片地址与衣物记录。","衣物出现在个人衣橱列表。"),
        ("UC-03","生成智能搭配","用户提交场景和约束，系统返回成套搭配方案。","用户已登录，衣橱至少有可选衣物。","解析场景；检索候选；规则过滤；请求模型；解析 JSON；校验单品归属；保存方案。","用户看到可解释方案，可收藏或替换单品。")], start=1):
        heading(doc,f"2.3.{index} {name}",3); caption(doc,f"表2.{index} {name}用例描述")
        add_table(doc,["内容","说明"],[["用例编号",code],["用例名称",name],["用例说明",desc],["前置条件",pre],["基本流程",flow],["异常路径","身份失效时返回 401；参数不完整时返回 400；模型服务不可用时触发降级提示。"],["后置条件",post]],[3.2,11])
    heading(doc,"2.4 技术方案",2)
    text(doc,"前端采用 Vue 3、Vite、TypeScript、Pinia 和 Element Plus；后端采用 Java 17、Spring Boot 3、Spring Security、MyBatis-Plus 与 Spring AI Alibaba。模型能力通过阿里云百炼通义千问 API 接入，天气能力通过 WeatherClient 调用开源 Open-Meteo API，并由服务层统一转换为业务天气对象；风潮数据通过独立的趋势来源适配器定时采集并归一化。数据层使用 PostgreSQL 16 和 pgvector 存储业务与向量数据，Redis 用于短期缓存和限流计数。部署时使用 Docker Compose 编排前端静态资源、后端服务、数据库和缓存。")
    heading(doc,"2.4.1 系统架构图",3)
    doc.add_picture(str(ASSETS/"architecture.png"),width=Cm(16)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; caption(doc,"图2.2 系统总体架构图")
    heading(doc,"2.4.2 数据库命名规则",3)
    text(doc,"数据库表名使用小写单数名词与下划线，例如 garment、outfit_plan；主键统一为 id；外键字段采用 xxx_id；创建和更新时间采用 created_at、updated_at；逻辑删除字段采用 is_deleted。业务表不保存明文密码、模型密钥或可直接识别用户的原始敏感信息。")
    heading(doc,"2.4.3 类命名规则",3)
    text(doc,"Java 类名使用 UpperCamelCase，例如 RecommendationService、GarmentController、OutfitPlanVO。实体、DTO、VO、Service 和 Mapper 使用后缀区分职责，避免以模糊缩写命名。")
    heading(doc,"2.4.4 方法、参数、成员变量、局部变量命名规则",3)
    text(doc,"方法名、参数名、成员变量和局部变量统一使用 lowerCamelCase，例如 generateOutfitPlan、userId、sceneRequest。布尔变量以 is、has 或 can 开头；不使用单字符业务变量，不在方法中混合数据库访问、模型调用和响应组装。")
    heading(doc,"2.4.5 包名结构",3)
    text(doc,"后端包结构划分为 com.fashion.recommendation.controller、service、service.impl、mapper、entity、dto、vo、config、security、ai 和 common。前端按 views、components、stores、api、types 和 utils 划分，页面不直接拼接接口地址。")
    heading(doc,"3 设计内容（过程）与说明",1)
    heading(doc,"3.1 数据库设计",2)
    text(doc,"数据库以“用户 - 衣物 - 推荐方案”为核心主线。sys_user 与 garment 为一对多关系；outfit_plan 记录一次完整推荐；outfit_item 表连接方案和衣物；recommendation_session 保存模型调用上下文摘要与状态。趋势聚合新增 trend_snapshot 表，保存 sourcePlatform、sourceUrl、title、topicTags、publishedAt、fetchedAt、heatScore、status 和 contentHash；以 sourcePlatform 与 sourceUrl 建立唯一约束，contentHash 用于去重。个人分析使用 user_style_profile 表保存 userId、profileVersion、styleTags、tryStyleTags、colorSuggestions、itemSuggestions、reasonSummary、modelName、generatedAt 和 sourceProfileHash；用户基础信息或衣橱标签变化后使旧档案失效。设计中将高频查询字段建立复合索引，并为用户私有数据增加 user_id 过滤条件。")
    heading(doc,"3.1.1 数据库物理模型",3)
    doc.add_picture(str(ASSETS/"er.png"),width=Cm(16)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; caption(doc,"图3.1 核心实体关系图")
    heading(doc,"3.1.2 数据表设计",3)
    for no,name,rows in [
        ("3.1","sys_user 用户表",[["id","bigint","否","用户主键"],["username","varchar(32)","否","唯一用户名"],["password_hash","varchar(100)","否","BCrypt 密文"],["created_at","timestamp","否","创建时间"]]),
        ("3.2","garment 衣物表",[["id","bigint","否","衣物主键"],["user_id","bigint","否","所属用户"],["category","varchar(32)","否","上装/下装等"],["color_tags","varchar(255)","是","颜色标签"],["season_tags","varchar(64)","是","季节标签"],["image_url","varchar(512)","否","私有对象地址"]]),
        ("3.3","outfit_plan 搭配方案表",[["id","bigint","否","方案主键"],["user_id","bigint","否","所属用户"],["scene","varchar(128)","否","通勤/约会等"],["reasoning_summary","text","否","简要推荐理由"],["status","varchar(16)","否","生成状态"],["created_at","timestamp","否","生成时间"]])]:
        text(doc,f"（{no[-1]}）{name}",False); caption(doc,f"表{no} {name}")
        add_table(doc,["字段名","类型","可空","说明"],rows,[3.2,3.2,2,5.8])
    heading(doc,"3.2 界面设计",2)
    text(doc,"“知己”前台计划实现登录注册页、搭配页、衣橱页、风潮页、方案档案页与设定页等页面。首页以钟表式功能切换替代常规顶部菜单：衣橱、风潮、搭配、档案和设定以短词固定分布，当前功能使用深色文字与两根指针表示；账号和通知等低频操作保留在右上角。")
    heading(doc,"3.2.1 用户登录页面",3)
    text(doc,"登录页面提供账号、密码和登录按钮，错误凭证只提示认证失败，不区分账号不存在或密码错误。登录成功后保存短期访问令牌，并跳转到首页。")
    heading(doc,"3.2.2 用户注册页面",3)
    text(doc,"注册页面要求填写用户名、密码与确认密码，并完成基础格式校验。后端对用户名设置唯一约束，对密码进行 BCrypt 散列后保存，禁止在前端或日志中记录明文。")
    heading(doc,"3.2.3 衣橱页面",3)
    text(doc,"衣橱页面按品类、颜色、季节和状态提供筛选，卡片展示衣物图片、标签和编辑入口。用户上传衣物后必须完成最少的品类与颜色标注，避免未经标注的数据直接进入推荐候选集。")
    heading(doc,"3.2.4 首页推荐与功能切换页面",3)
    doc.add_picture(str(ASSETS/"ui.png"),width=Cm(16)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; caption(doc,"图3.2 首页放射式功能切换原型")
    heading(doc,"3.2.5 档案页面",3); text(doc,"档案页展示已保存的搭配方案，支持按时间和场景筛选，并允许复用场景重新生成。页面不展示模型内部思维链，只展示面向用户的简短理由与规则命中说明。")
    heading(doc,"3.2.6 设定页面",3); text(doc,"设定页用于维护偏好、常用场景和隐私选项；用户可管理收藏反馈，并随时删除自己的反馈数据。")
    heading(doc,"3.2.7 风潮趋势页面",3)
    text(doc,"风潮页面分为“全局风潮”和“我的风格”两个区域。全局风潮以抖音趋势为主，按来源平台和主题标签展示趋势卡片，卡片包括标题或话题、热度、来源、发布时间或采集时间和原文跳转链接；页面必须标注“聚合自公开来源”，不复制受限全文、不展示个人账号资料，也不将其他平台内容伪装为本站原创。我的风格区域读取已保存的 LLM 风格档案，展示当前推荐风格、推荐搭配、建议单品、颜色搭配和可尝试风格。某一来源刷新失败时，保留最后一次成功数据并显示其采集时间。")
    heading(doc,"3.3 类设计",2)
    text(doc,"后端按控制层、应用服务层、领域对象与基础设施层分离。RecommendationService 负责协调场景解析、候选检索、规则校验和模型调用；GarmentService 只负责衣物业务；AiClientAdapter 隔离具体模型供应商差异。这样的划分避免让控制器承担业务流程，也使模型服务可替换。")
    heading(doc,"3.3.1 衣橱管理业务类设计",3)
    text(doc,"衣橱管理由 GarmentController、GarmentService、GarmentMapper、GarmentEntity 与 GarmentVO 构成。服务层完成归属校验、标签规范化和图片地址保存，Mapper 仅负责持久化访问。")
    heading(doc,"3.3.2 智能推荐业务类设计",3)
    text(doc,"智能推荐由 RecommendationController、RecommendationService、WardrobeRetriever、OutfitRuleValidator 和 AiClientAdapter 协作完成。生成前必须取得候选衣物，生成后必须再次校验衣物 ID 与用户归属。")
    caption(doc,"表3.4 核心类职责说明")
    add_table(doc,["类/组件","职责","关键约束"],[
        ["RecommendationController","接收推荐请求并返回统一响应","不直接访问数据库或拼接提示词。"],["RecommendationService","编排推荐流程与结果保存","必须校验模型返回的衣物 ID 属于当前用户。"],["WardrobeRetriever","按场景筛选衣物并召回相似标签","空衣橱时返回明确引导。"],["OutfitRuleValidator","校验季节、品类和可穿状态","规则失败不展示为可执行方案。"],["AiClientAdapter","封装百炼通义千问 API","超时、限流与异常统一转换为领域错误。"],["TrendAggregationService","调度来源适配器、归一化、去重并写入趋势快照","抖音适配器为主，不抓取登录态或受限内容，所有展示项必须可追溯至 sourceUrl。"],["PersonalStyleProfileService","根据用户基础信息和手动衣橱标签生成、缓存风格档案","仅输出固定 Schema；输入变化后使旧档案失效。"]],[4,5.2,5])
    heading(doc,"3.3.3 天气数据服务设计",3)
    text(doc,"天气数据使用开源项目 Open-Meteo（https://github.com/open-meteo/open-meteo）的公开 API。毕业设计为非商业使用，不需要配置 API Key；项目不复制、修改或自建其服务端。后端不得让前端直接调用第三方接口，而是由 WeatherClient 集中处理地址、超时、字段映射和错误转换，避免第三方响应结构泄漏到业务层。")
    caption(doc,"表3.5 天气服务调用契约")
    add_table(doc,["步骤","第三方地址","入参","处理结果"],[
        ["城市解析","https://geocoding-api.open-meteo.com/v1/search","name, count=1, language=zh, format=json","取首个结果的 latitude、longitude；无结果返回 WEATHER_LOCATION_NOT_FOUND。"],
        ["天气查询","https://api.open-meteo.com/v1/forecast","latitude, longitude, current=temperature_2m,apparent_temperature,precipitation,weather_code,wind_speed_10m, timezone=Asia/Shanghai","映射为温度、体感温度、降水、天气现象和风速。"]],[2,5.6,4.3,7.3])
    text(doc,"建议定义 WeatherSnapshotVO，字段为 city、observedAt、temperatureC、apparentTemperatureC、precipitationMm、weatherCode、windSpeedKmh 和 source。weatherCode 保留 Open-Meteo 的 WMO 数值供规则层判断；前端展示文案由后端字典转换，不能让模型自行猜测天气含义。搭配规则至少使用体感温度、降水和风速：低温提高保暖层级，降水提示防水外层或雨具，较大风速提示防风外套。")
    text(doc,"查询成功结果缓存 15 分钟，缓存键包含标准化城市名和经纬度。连接与读取超时均设置为 3 秒，网络异常、429 或 5xx 时重试一次；仍失败则返回 WEATHER_SERVICE_UNAVAILABLE，不将旧天气伪装成实时数据。用户已手动填写温度或天气时，可继续按手动数据生成搭配，但响应中需标明 weatherSource=MANUAL。")
    heading(doc,"3.3.4 风潮趋势聚合设计",3)
    text(doc,"TrendAggregationService 按固定周期调用各 TrendSourceAdapter。抖音适配器为主来源，默认每 3 小时刷新一次；其他启用来源默认每 6 小时刷新一次。每个适配器只负责一个来源，使用官方开放接口、许可的 RSS/公开页面或已授权数据源获取数据；HTTP 请求设置来源级超时和速率限制，禁止使用账号密码、验证码识别、代理轮换或其他绕过平台限制的手段。获取结果统一转换为 TrendSnapshot，再按 sourcePlatform + sourceUrl 去重并保存。")
    text(doc,"建议为每个来源单独配置启用开关、刷新周期、最大条数和最后成功时间。Redis 缓存风潮首页查询结果；PostgreSQL 保存可展示快照。数据源异常时，接口返回现有快照和 stale=true，不因单个平台失败而使风潮页面整体不可用。采集日志只记录来源、时间、HTTP 状态和条数，不记录用户个人资料或完整正文。")
    heading(doc,"3.3.5 个人风格档案设计",3)
    text(doc,"PersonalStyleProfileService 在用户首次完善基础信息或主动刷新时调用百炼通义千问 API。输入仅包含用户主动填写的风格偏好、色彩倾向、常用场合、预算敏感度和手动衣橱标签；输出必须满足 StyleProfileSchema，包含 styleTags、tryStyleTags、colorSuggestions、itemSuggestions 和 reasonSummary。模型不得编造用户未提供的体型、身份或消费能力，也不得将风潮热度当作用户偏好。")
    text(doc,"用户资料和衣橱标签变化后，将 sourceProfileHash 与已保存档案比对；不一致时标记档案过期。页面读取已保存档案，不在每次进入风潮页时调用模型。模型调用超时或返回非法 JSON 时保留上一份有效档案并标记 styleProfileStale=true；首次生成失败则返回可理解的补充信息提示。")
    heading(doc,"3.3.6 全系统后端代码包图",3)
    text(doc,"系统代码包以 controller 接收 HTTP 请求，service 承担业务编排，mapper 访问 PostgreSQL，ai 包封装模型能力，security 包维护认证授权。包之间通过接口和 DTO/VO 传递数据，避免前端对象直接映射数据库实体。")
    heading(doc,"3.4 顺序图设计",2)
    text(doc,"生成推荐时，前端先提交场景与用户约束，后端完成身份校验；检索组件获取候选衣物，规则组件剔除不可用候选；模型适配器请求 LLM 返回 JSON；服务层对 JSON 做结构校验和所有权校验，成功后持久化方案并返回前端。任何一个外部调用失败均不得写入“已完成”的方案记录。")
    heading(doc,"3.4.1 用户登录认证顺序图",3)
    text(doc,"前端提交账号和密码，AuthController 调用 AuthService 校验密码散列，认证成功后签发 JWT 并返回；失败时不生成令牌，也不暴露用户是否存在。")
    heading(doc,"3.4.2 新增衣物顺序图",3)
    text(doc,"前端上传图片和衣物标签，GarmentController 校验令牌和文件类型，GarmentService 保存对象地址及元数据，随后返回衣物详情。重复提交由前端禁用按钮和后端请求幂等键共同控制。")
    heading(doc,"3.4.3 智能推荐顺序图",3)
    caption(doc,"表3.6 智能推荐顺序说明")
    add_table(doc,["顺序","参与者","操作"],[
        ["1","Vue 前端 -> RecommendationController","提交 scene、temperature、styleTags 和约束。"],["2","Controller -> RecommendationService","校验 JWT 与请求参数。"],["3","Service -> WardrobeRetriever","检索当前用户可穿衣物。"],["4","Service -> OutfitRuleValidator","过滤季节、品类冲突和不可穿状态。"],["5","Service -> AiClientAdapter","发送受约束的候选集与 JSON Schema。"],["6","Service -> Repository","校验并保存方案与明细。"],["7","Service -> Vue 前端","返回方案、理由和可替换单品。"]],[1.2,4.6,8.4])
    heading(doc,"3.4.4 收藏与反馈顺序图",3)
    text(doc,"用户对方案点击收藏或评分，前端提交 outfitId 与评价，后端先校验方案归属，再写入 favorite_outfit 记录。若方案不存在或属于其他用户，接口返回无权访问，不能以评分接口探测他人数据。")
    heading(doc,"3.5 API 接口设计",2)
    text(doc,"接口遵循 REST 风格，统一前缀为 /api/v1，使用 HTTPS 传输。成功响应返回 code、message 和 data；业务异常使用可枚举错误码；文件上传接口限制 MIME 类型、尺寸和鉴权。")
    caption(doc,"表3.7 核心接口描述")
    add_table(doc,["接口","方法","地址","关键入参","说明"],[
        ["登录","POST","/auth/login","username, password","校验账号并返回 JWT。"],["新增衣物","POST","/garments","image, category, colorTags","保存用户衣物及标签。"],["查询衣橱","GET","/garments","category, season","仅返回当前用户数据。"],["查询天气","GET","/weather/current","city","后端调用 Open-Meteo，返回 WeatherSnapshotVO。"],["查询风潮","GET","/trends","platform, topic, cursor","返回已归一化的趋势快照和来源链接。"],["查询个人风格","GET","/me/style-profile","无","返回已保存的风格档案及是否过期。"],["刷新个人风格","POST","/me/style-profile/refresh","无","调用百炼 API 生成新档案。"],["生成推荐","POST","/recommendations","scene, weather, styleTags","生成并保存结构化方案。"],["收藏方案","POST","/outfits/{id}/favorite","rating, comment","保存个人反馈。"]],[2.4,1.5,3.4,4.3,3.8])
    heading(doc,"3.5.1 用户登录与注册模块",3)
    text(doc,"登录接口为 POST /api/v1/auth/login，注册接口为 POST /api/v1/auth/register。两者均接受 JSON 格式参数，密码只可通过 HTTPS 传输，响应中不得返回 passwordHash。")
    heading(doc,"3.5.2 我的衣橱管理模块",3)
    text(doc,"衣物新增、修改、删除和查询接口统一使用 /api/v1/garments。所有查询条件都自动附加当前 userId，删除为逻辑删除并使该衣物退出新的推荐候选集。")
    heading(doc,"3.5.3 智能推荐与反馈模块",3)
    text(doc,"推荐接口为 POST /api/v1/recommendations，收藏与反馈接口为 POST /api/v1/outfits/{id}/favorite。推荐结果使用固定 JSON 结构返回，外部模型异常统一映射为可识别业务错误。")
    heading(doc,"3.5.4 天气查询模块",3)
    text(doc,"天气查询接口为 GET /api/v1/weather/current?city={city}。响应 data 使用 WeatherSnapshotVO；city 为空时返回 400，城市无法解析时返回 404 和 WEATHER_LOCATION_NOT_FOUND，第三方服务不可用时返回 503 和 WEATHER_SERVICE_UNAVAILABLE。Open-Meteo 未使用 API Key，因此不得在 .env、日志或接口响应中添加虚构的天气密钥字段。")
    heading(doc,"3.5.5 风潮趋势模块",3)
    text(doc,"趋势查询接口为 GET /api/v1/trends?platform={platform}&topic={topic}&cursor={cursor}。响应项包含 platform、title、topicTags、heatScore、publishedAt、fetchedAt、sourceUrl 和 stale，不返回受限全文或个人资料。抖音为默认 platform；platform 仅接受已配置来源。未配置、被暂停或无可用快照的来源不触发临时爬取，而是返回空列表和来源状态，避免用户请求放大对第三方的访问压力。")
    heading(doc,"3.5.6 个人风格档案模块",3)
    text(doc,"GET /api/v1/me/style-profile 返回当前登录用户的已保存档案；POST /api/v1/me/style-profile/refresh 仅在用户主动触发时调用百炼 API。响应包含 styleTags、tryStyleTags、colorSuggestions、itemSuggestions、reasonSummary、generatedAt 和 stale。用户资料不完整时返回 422 和 STYLE_PROFILE_INCOMPLETE；百炼服务异常时首次生成返回 503，已有有效档案则返回 200 且 stale=true。")
    heading(doc,"3.6 项目测试",2)
    text(doc,"测试围绕真实业务约束设计。以下表格给出测试计划与预期结果，不将未运行的测试标记为通过。最终提交前需在实际部署环境中补充执行日期、测试人员、接口响应截图和自动化测试报告。")
    heading(doc,"3.6.1 单元测试计划",3)
    caption(doc,"表3.8 单元测试计划表")
    add_table(doc,["模块","测试点","输入或条件","预期结果","实测状态"],[
        ["认证服务","密码校验","正确/错误密码","正确密码生成令牌，错误密码返回认证失败。","待实测"],["衣橱服务","用户数据隔离","访问其他用户 garmentId","返回无权限或不存在，不泄露数据。","待实测"],["规则校验","季节冲突","夏季场景 + 厚羽绒服","候选被过滤，不进入模型提示词。","待实测"],["天气服务","天气字段映射","Open-Meteo 返回天气码和温度","正确生成 WeatherSnapshotVO 与展示文案。","待实测"],["个人风格服务","资料变更失效","修改色彩倾向或衣橱标签","旧档案标记过期，不直接复用。","待实测"],["模型解析","非法 JSON","百炼返回缺字段文本","解析失败后降级，不保存成功档案。","待实测"]],[2.3,3.4,4.2,4.9,2])
    heading(doc,"3.6.2 接口测试计划",3)
    caption(doc,"表3.9 接口测试计划表")
    add_table(doc,["接口","测试场景","预期 HTTP 结果","实测状态"],[
        ["POST /auth/login","密码错误","401，返回统一错误码","待实测"],["POST /garments","上传非图片或超尺寸文件","400，拒绝保存文件","待实测"],["GET /weather/current","不存在的城市","404，返回 WEATHER_LOCATION_NOT_FOUND","待实测"],["GET /weather/current","Open-Meteo 超时或 5xx","503，不返回伪造实时天气","待实测"],["GET /trends","单一来源刷新失败但存在快照","200，返回快照且 stale=true","待实测"],["POST /me/style-profile/refresh","基础资料不完整","422，返回 STYLE_PROFILE_INCOMPLETE","待实测"],["POST /recommendations","衣橱为空","200 或 422，返回补充衣物引导","待实测"],["POST /recommendations","模型接口超时","503 或可重试提示，不产生成功方案","待实测"]],[3.6,4.6,5.2,2])
    heading(doc,"3.6.3 集成测试计划",3)
    caption(doc,"表3.10 核心业务链路测试计划表")
    add_table(doc,["链路","步骤","预期结果","实测状态"],[
        ["新用户首次推荐","注册 -> 上传 3 件衣物 -> 选择通勤场景 -> 生成方案","方案中的所有衣物均属于该用户，且包含理由。","待实测"],["快速重复提交","连续两次点击生成","后端幂等或前端禁用，避免重复保存相同会话。","待实测"],["跨用户越权","用户 A 请求用户 B 的方案 ID","拒绝访问，不返回方案详情。","待实测"],["趋势来源降级","抖音来源返回 5xx，查询风潮页","其他来源正常返回，抖音仅展示历史快照和采集时间。","待实测"],["个人风格降级","百炼调用超时且已有档案","返回上一份有效档案并标记 stale=true。","待实测"],["外部服务故障","模拟 LLM 5xx","用户看到降级说明，系统日志记录关联请求号。","待实测"]],[3,5.2,5.2,2])
    heading(doc,"3.6.4 单元测试报告",3)
    text(doc,"当前文档仅提供单元测试报告填写框架。实现完成后应补充 JUnit 5 或 Mockito 的实际执行日期、测试类、断言数量、覆盖率和失败项；未执行前的状态统一为“待实测”。")
    caption(doc,"表3.11 单元测试报告")
    add_table(doc,["模块","测试类","预期覆盖点","实际结果","状态"],[["认证模块","AuthServiceTest","密码匹配、令牌签发、错误密码","待填入真实结果","待实测"],["衣橱模块","GarmentServiceTest","归属校验、标签校验、逻辑删除","待填入真实结果","待实测"],["推荐模块","RecommendationServiceTest","规则过滤、非法 JSON 降级","待填入真实结果","待实测"]],[2.4,3.6,4.5,3.6,2])
    heading(doc,"3.6.5 接口测试报告",3)
    text(doc,"接口报告应由 Postman、Apifox 或自动化脚本导出后填写。重点核对 HTTP 状态码、错误码、响应字段、鉴权边界与服务超时行为。")
    caption(doc,"表3.12 接口测试报告")
    add_table(doc,["接口","测试数据","预期","实际","状态"],[["POST /auth/login","错误密码","401 与统一错误码","待填入真实响应","待实测"],["POST /garments","非图片文件","400，文件不落库","待填入真实响应","待实测"],["POST /recommendations","空衣橱","补充衣物引导","待填入真实响应","待实测"]],[3.3,3.1,3.3,3.3,2])
    heading(doc,"3.6.6 集成测试报告",3)
    text(doc,"集成测试以“登录 - 上传衣物 - 生成方案 - 收藏反馈”为核心业务链路。提交前应在与目标部署一致的环境中执行，并附上运行日志、接口截图和缺陷修复记录。")
    caption(doc,"表3.13 集成测试报告")
    add_table(doc,["链路","预期结果","实际结果","状态"],[["登录至首页","身份认证后正确进入首页","待填入真实结果","待实测"],["上传至推荐","方案仅引用当前用户可穿衣物","待填入真实结果","待实测"],["推荐至反馈","收藏后历史方案可查询","待填入真实结果","待实测"]],[3.4,4.5,4.5,2])
    heading(doc,"4 设计总结",1)
    heading(doc,"4.1 部署手册",2)
    text(doc,"系统部署到云服务器，采用 Docker Compose 编排。部署前需准备域名或服务器 IP、HTTPS 证书、PostgreSQL 数据卷、Redis 数据卷、对象存储配置和阿里云百炼 API Key。所有密钥通过环境变量注入，不提交到代码仓库；生产 Redis 使用受支持版本，不直接复用开发机的 Redis 3.2.100。")
    heading(doc,"4.1.1 初始化数据服务",3)
    bullet(doc,"创建 PostgreSQL 16 数据库 fashion_recommendation，并执行 db/migration 下的结构脚本。")
    bullet(doc,"启用 pgvector 扩展，创建衣物向量索引；启动 Redis 7，设置访问密码。")
    bullet(doc,"创建对象存储桶 garments-private，并配置仅由后端签名访问。")
    heading(doc,"4.1.2 部署应用并启动服务",3)
    bullet(doc,"在服务器目录中配置 .env：DB_URL、DB_USERNAME、DB_PASSWORD、REDIS_PASSWORD、DASHSCOPE_API_KEY、BAILIAN_MODEL、OBJECT_STORAGE_*。Open-Meteo 默认调用不需要 WEATHER_API_KEY；如需替换服务，只新增 WEATHER_BASE_URL 和 WeatherClient 配置，不改动推荐业务类。趋势来源按平台单独配置启用状态、允许的 base URL、刷新周期和速率限制；抖音默认 3 小时、其他来源默认 6 小时，禁止配置用户账号、密码或绕过访问限制的参数。")
    bullet(doc,"执行 docker compose up -d --build，检查 frontend、backend、postgres、redis 容器均为 healthy。")
    bullet(doc,"访问 /actuator/health 验证后端健康状态；未配置 LLM Key 时推荐接口应明确提示服务未配置。")
    heading(doc,"4.1.3 访问程序",3)
    text(doc,"在可访问服务器的浏览器中打开 https://{domain}/。首次使用时注册账号，上传至少三件已标注衣物，再进入“开始搭配”提交场景。正式部署必须启用 HTTPS、反向代理限流和数据库定期备份。")
    heading(doc,"4.2 用户操作手册",2)
    heading(doc,"4.2.1 用户登录",3); text(doc,"打开系统首页后输入已注册账号和密码，点击登录。登录失败时按提示检查凭证；连续失败保护策略应由实际部署配置决定。")
    heading(doc,"4.2.2 用户注册",3); text(doc,"没有账号的用户进入注册页面，填写用户名、密码和确认密码。注册成功后可返回登录页进行身份认证；用户名重复时系统应给出明确提示。")
    heading(doc,"4.2.3 新建衣橱",3); text(doc,"登录后进入“我的衣橱”，点击新增衣物，上传清晰图片并填写品类、颜色和适用季节。只有状态为“可穿”的衣物会进入推荐候选。")
    heading(doc,"4.2.4 获取智能搭配",3); text(doc,"在首页选择“开始搭配”，填写通勤、约会、旅行等场景，可补充温度、风格和禁忌。点击生成后查看推荐单品及理由；若衣橱缺少关键单品，系统应提示补充而不是虚构已有衣物。")
    heading(doc,"4.2.5 收藏与反馈",3); text(doc,"对满意方案点击收藏，可对方案评分或添加简短反馈。用户可在“历史方案”中查看之前的推荐并复用场景。删除衣物后，历史方案保留文字记录，但不再将该衣物作为新的推荐候选。")
    heading(doc,"4.3 致谢",2)
    text(doc,"本毕业设计的完成离不开指导教师在选题、技术方案和文档规范方面的指导，也感谢在需求分析、测试反馈和资料查阅过程中提供帮助的老师与同学。通过本项目，进一步理解了前后端分离架构、数据建模、AI 服务编排和软件测试中“可验证性优先”的工程原则。")
    heading(doc,"参考文献",1)
    refs=[
        "[1] Spring. Spring Boot Reference Documentation 3.3 [EB/OL]. 2024.",
        "[2] Vue.js Team. Vue.js Guide: Introduction [EB/OL]. 2025.",
        "[3] PostgreSQL Global Development Group. PostgreSQL 16 Documentation [EB/OL]. 2023.",
        "[4] Bai J, Bai S, Chu Y, et al. Qwen Technical Report [EB/OL]. 2023.",
        "[5] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks [C]. NeurIPS, 2020.",
        "[6] OWASP Foundation. Application Security Verification Standard 4.0.3 [EB/OL]. 2021.",
        "[7] MyBatis-Plus. MyBatis-Plus Documentation [EB/OL]. 2024.",
        "[8] Redis Ltd. Redis Documentation [EB/OL]. 2024."
    ]
    for ref in refs: text(doc,ref,False)
    heading(doc,"附 录",1)
    heading(doc,"附录 A 推荐结果 JSON 示例",2)
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.74); p.paragraph_format.line_spacing=1.2
    r=p.add_run('{\n  "scene": "通勤",\n  "outfitItems": [{"garmentId": 101, "role": "top"}],\n  "reason": "浅色上装与深色下装形成通勤场景所需的稳定对比。",\n  "alternatives": []\n}')
    set_font(r,"Consolas",9)
    heading(doc,"附录 B 数据与模型使用边界",2)
    text(doc,"系统仅将用户主动提交的衣物标签、场景和偏好用于当前推荐服务。生产环境中应取得用户授权、提供数据删除入口，并对第三方模型接口的请求内容执行最小化传输。模型输出仅作为搭配建议，不构成医疗、消费或身份判断结论。")
    settings=doc.settings.element
    update=OxmlElement("w:updateFields"); update.set(qn("w:val"),"true"); settings.append(update)
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__": build()
